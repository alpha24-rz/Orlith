import logging
import asyncio
import pdfplumber
import docx
import pypdfium2 as pdfium
import numpy as np
import torch
import easyocr
import re
from typing import Tuple, List, Dict
from models.document import Document
from services.context import PipelineContext
import hashlib

logger = logging.getLogger(__name__)

_easyocr_reader = None

def get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        cuda_available = torch.cuda.is_available()
        logger.info(f"Initializing EasyOCR reader (en, id) on {'GPU (CUDA)' if cuda_available else 'CPU'}...")
        _easyocr_reader = easyocr.Reader(['en', 'id'], gpu=cuda_available)
    return _easyocr_reader

class TextExtractionService:
    async def extract(self, document: Document, ocr: bool, ctx: PipelineContext) -> Tuple[List[Dict], Dict]:
        ctx.transition("extracting")
        start_time = asyncio.get_event_loop().time()
        
        pages_data, metadata = await asyncio.to_thread(
            self._extract_sync, document.file_path, document.file_type, ocr
        )
        
        # Calculate text hash and update metadata
        full_text = " ".join([p["text"] for p in pages_data])
        normalized_text = re.sub(r'\s+', ' ', full_text).strip()
        metadata["text_hash"] = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
        
        # Heuristic Metadata enrichment (P7)
        metadata["title"] = self._extract_title(pages_data)
        metadata["author"] = self._extract_author(full_text)
        
        # P6: Non-blocking Language Detection
        try:
            from langdetect import detect
            metadata["language"] = detect(normalized_text[:1000]) if normalized_text else "unknown"
        except Exception as e:
            logger.warning(f"Language detection failed: {e}")
            metadata["language"] = "unknown"
        
        elapsed = asyncio.get_event_loop().time() - start_time
        ctx.record("extraction_time", elapsed)
        
        return pages_data, metadata

    def _extract_sync(self, file_path: str, file_type: str, ocr: bool) -> Tuple[List[Dict], Dict]:
        file_path_lower = file_path.lower()
        metadata = {"ocr_applied": False, "page_count": 1}
        pages_data = []

        if file_path_lower.endswith(".txt") or file_path_lower.endswith(".md") or file_type in ("text/plain", "text/markdown"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text_content = f.read()
            pages_data.append({"text": text_content, "page_number": 1})
            metadata["page_count"] = 1

        elif file_path_lower.endswith(".docx") or file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(file_path)
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text:
                        text_parts.append(paragraph.text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)

                text_content = "\n".join(text_parts)
                pages_data.append({"text": text_content, "page_number": 1})
                metadata["page_count"] = len(doc.paragraphs)
            except Exception as docx_err:
                logger.warning(f"python-docx parsing failed: {docx_err}. Attempting raw XML zip fallback...")
                try:
                    import zipfile
                    import xml.etree.ElementTree as ET
                    with zipfile.ZipFile(file_path) as docx_zip:
                        xml_content = docx_zip.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    text_parts = []
                    for elem in root.iter():
                        if elem.tag.endswith('}t') and elem.text:
                            text_parts.append(elem.text)
                    text_content = " ".join(text_parts)
                    if not text_content.strip():
                        raise ValueError("No text extracted from word/document.xml")
                    pages_data.append({"text": text_content, "page_number": 1})
                    metadata["page_count"] = 1
                    logger.info("Raw XML zip fallback successfully extracted text from DOCX.")
                except Exception as fallback_err:
                    raise RuntimeError(f"Word (DOCX) text extraction failed. python-docx error: {docx_err}. XML zip fallback error: {fallback_err}")

        elif file_path_lower.endswith(".pdf") or file_type == "application/pdf":
            page_count = 0
            total_text = ""
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            pages_data.append({"text": page_text, "page_number": i + 1})
                            total_text += page_text
            except Exception as e:
                logger.warning(f"pdfplumber text extraction failed, falling back to OCR if available. Error: {e}")

            metadata["page_count"] = page_count

            if ocr or len(total_text.strip()) < 50:
                logger.info("PDF has very little text or OCR was explicitly requested. Attempting OCR...")
                try:
                    pages_data.clear()
                    pdf_doc = pdfium.PdfDocument(file_path)
                    page_count = len(pdf_doc)
                    metadata["page_count"] = page_count

                    reader = get_easyocr_reader()

                    for i, page in enumerate(pdf_doc):
                        image = page.render(scale=2).to_pil()
                        image_np = np.array(image)
                        results = reader.readtext(image_np, detail=0)
                        page_text = "\n".join(results)
                        if page_text:
                            pages_data.append({"text": page_text, "page_number": i + 1})

                    metadata["ocr_applied"] = True
                except Exception as e:
                    raise RuntimeError(f"OCR processing failed: {str(e)}")
        else:
            raise ValueError(f"Unsupported file format: {file_type} or file extension for {file_path}")

        return pages_data, metadata

    def _extract_title(self, pages: List[Dict]) -> Dict:
        """Heuristic title extraction."""
        if not pages:
            return {"value": "unknown", "confidence": 0.0}
        
        first_page = pages[0]["text"].split("\n")
        # Assume first non-empty line could be title
        for line in first_page:
            clean_line = line.strip()
            if len(clean_line) > 5 and len(clean_line) < 100:
                return {"value": clean_line, "confidence": 0.75}
                
        return {"value": "unknown", "confidence": 0.0}
        
    def _extract_author(self, full_text: str) -> Dict:
        """Heuristic author extraction."""
        match = re.search(r'(?i)(?:author|by|penulis)\s*[:]\s*([a-zA-Z\s,]+)', full_text[:2000])
        if match:
            author = match.group(1).strip()
            if len(author) > 2 and len(author) < 100:
                return {"value": author, "confidence": 0.85}
        return {"value": "unknown", "confidence": 0.0}
