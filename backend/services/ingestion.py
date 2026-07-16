from models import Document, Workspace
from core.chroma import get_workspace_collection
from services.ai.gateway import LLMGateway
from sqlalchemy.ext.asyncio import AsyncSession
from langchain_text_splitters import RecursiveCharacterTextSplitter
from core.config import settings
from core.websocket import manager
import uuid
import os
import hashlib
import pdfplumber
import docx
import pypdfium2 as pdfium
from PIL import Image
import logging
import easyocr
import numpy as np
import asyncio
import torch

logger = logging.getLogger(__name__)

_ingestion_semaphore = asyncio.Semaphore(settings.INGESTION_CONCURRENCY)

_easyocr_reader = None

def get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        cuda_available = torch.cuda.is_available()
        logger.info(f"Initializing EasyOCR reader (en, id) on {'GPU (CUDA)' if cuda_available else 'CPU'}...")
        _easyocr_reader = easyocr.Reader(['en', 'id'], gpu=cuda_available)
    return _easyocr_reader


def get_file_hash(file_path: str) -> str:
    """Compute SHA-256 hash of a file."""
    hasher = hashlib.sha256()
    with open(file_path, "rb") as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()


def extract_text_from_file(
    file_path: str, file_type: str, ocr: bool = False
) -> tuple[list[dict], dict]:
    """
    Extracts text from various file formats.
    Returns: (list_of_pages, metadata_dict)
    where list_of_pages = [{"text": str, "page_number": int}]
    """
    file_path_lower = file_path.lower()
    metadata = {"ocr_applied": False, "page_count": 1}
    pages_data = []

    # TXT / MD Files
    if (
        file_path_lower.endswith(".txt")
        or file_path_lower.endswith(".md")
        or file_type in ("text/plain", "text/markdown")
    ):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text_content = f.read()
        pages_data.append({"text": text_content, "page_number": 1})
        metadata["page_count"] = 1

    # Word (DOCX) Files
    elif (
        file_path_lower.endswith(".docx")
        or file_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        try:
            doc = docx.Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            text_content = "\n".join(text_parts)
            pages_data.append({"text": text_content, "page_number": 1})
            metadata["page_count"] = len(doc.paragraphs)
        except Exception as docx_err:
            logger.warning(f"python-docx parsing failed: {docx_err}. Attempting raw XML zip fallback...")
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(file_path) as docx_zip:
                    xml_content = docx_zip.read('word/document.xml')
                
                root = ET.fromstring(xml_content)
                text_parts = []
                for elem in root.iter():
                    if elem.tag.endswith('}t') and elem.text:
                        text_parts.append(elem.text)
                
                text_content = " ".join(text_parts)
                if not text_content.strip():
                    raise ValueError("No text extracted from word/document.xml")
                
                pages_data.append({"text": text_content, "page_number": 1})
                metadata["page_count"] = 1
                logger.info("Raw XML zip fallback successfully extracted text from DOCX.")
            except Exception as fallback_err:
                raise RuntimeError(
                    f"Word (DOCX) text extraction failed. python-docx error: {docx_err}. "
                    f"XML zip fallback error: {fallback_err}"
                )

    # PDF Files
    elif file_path_lower.endswith(".pdf") or file_type == "application/pdf":
        page_count = 0
        total_text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_data.append({"text": page_text, "page_number": i + 1})
                        total_text += page_text
        except Exception as e:
            logger.warning(
                f"pdfplumber text extraction failed, falling back to OCR if available. Error: {e}"
            )

        metadata["page_count"] = page_count

        # Check if we should apply OCR
        if ocr or len(total_text.strip()) < 50:
            logger.info(
                "PDF has very little text or OCR was explicitly requested. Attempting OCR..."
            )

            try:
                pages_data.clear()  # clear whatever we got from pdfplumber
                pdf_doc = pdfium.PdfDocument(file_path)
                page_count = len(pdf_doc)
                metadata["page_count"] = page_count

                reader = get_easyocr_reader()

                for i, page in enumerate(pdf_doc):
                    image = page.render(scale=2).to_pil()
                    image_np = np.array(image)
                    results = reader.readtext(image_np, detail=0)
                    page_text = "\n".join(results)
                    if page_text:
                        pages_data.append({"text": page_text, "page_number": i + 1})

                metadata["ocr_applied"] = True
            except Exception as e:
                raise RuntimeError(f"OCR processing failed: {str(e)}")

    else:
        raise ValueError(
            f"Unsupported file format: {file_type} or file extension for {file_path}"
        )

    return pages_data, metadata


async def broadcast_status(document: Document, db: AsyncSession, error_msg: str = None):
    """Broadcast status updates to the workspace's websocket connections and create global notifications."""
    payload = {
        "event": "document_status",
        "data": {
            "id": document.id,
            "workspace_id": document.workspace_id,
            "filename": document.filename,
            "status": document.status,
            "file_size": document.file_size,
            "created_at": document.created_at.isoformat()
            if document.created_at
            else None,
            "content_hash": document.content_hash,
            "metadata": document.metadata_json,
            "error": error_msg,
        },
    }
    await manager.broadcast_to_workspace(document.workspace_id, payload)

    # Global notifications for final states
    if document.status in ("ready", "error", "failed"):
        from models import Workspace, WorkspaceMember, Notification
        from sqlalchemy import select
        
        workspace = await db.get(Workspace, document.workspace_id)
        if not workspace:
            return

        result = await db.execute(select(WorkspaceMember.user_id).where(WorkspaceMember.workspace_id == workspace.id))
        user_ids = set(result.scalars().all())
        user_ids.add(workspace.owner_id)

        title = "Extraction Completed" if document.status == "ready" else "Extraction Failed"
        desc = f"{document.filename} is ready." if document.status == "ready" else f"{document.filename}: {error_msg or 'Failed'}"
        ntype = "success" if document.status == "ready" else "error"

        notifications = []
        for uid in user_ids:
            notif = Notification(
                user_id=uid,
                workspace_id=workspace.id,
                title=title,
                description=desc,
                type=ntype
            )
            db.add(notif)
            notifications.append(notif)
        
        await db.commit()

        # broadcast to global channels
        for notif in notifications:
            payload_global = {
                "event": "notification",
                "data": {
                    "id": notif.id,
                    "title": notif.title,
                    "description": notif.description,
                    "type": notif.type,
                    "workspace_id": notif.workspace_id,
                    "created_at": notif.created_at.isoformat() if notif.created_at else None,
                    "is_read": False
                }
            }
            await manager.broadcast_to_workspace(f"global_{notif.user_id}", payload_global)


from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception
from httpx import HTTPStatusError, ConnectError, TimeoutException

def is_transient_error(e):
    if isinstance(e, HTTPStatusError):
        return e.response.status_code in (429, 502, 503, 504)
    return isinstance(e, (ConnectError, TimeoutException, ConnectionError))

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=8),
    retry=retry_if_exception(is_transient_error),
    before_sleep=lambda rs: logger.warning(f"Transient error, retrying embedding #{rs.attempt_number}...")
)
async def embed_document(
    document: Document, chunks: list, embedding_provider, model_name: str, db: AsyncSession, provider_name: str
):
    """Embeds chunks and stores them into ChromaDB and PostgreSQL Chunk tables."""
    if not chunks:
        raise ValueError("No chunks to embed.")

    from models.chunk import Chunk, ChunkEmbedding
    from sqlalchemy import delete

    # 1. Clean existing PostgreSQL chunks for this document
    await db.execute(delete(Chunk).where(Chunk.document_id == document.id))
    # ChunkEmbedding cascades on delete, so no need to delete them explicitly
    await db.commit()

    collection = get_workspace_collection(document.workspace_id)
    # Check if there are existing embeddings for this document (e.g., in a re-embed scenario)
    try:
        await asyncio.to_thread(collection.delete, where={"document_id": document.id})
    except Exception:
        pass  # Ignore if not found

    texts_to_embed = [c["text"] for c in chunks]
    embeddings = await embedding_provider.embed(texts_to_embed, model_name)

    # Prepare for Chroma
    ids = [f"{document.id}_{c['chunk_index']}" for c in chunks]
    metadatas = [
        {
            "document_id": document.id,
            "filename": document.filename,
            "page_number": c["page_number"],
            "chunk_index": c["chunk_index"],
            "parent_content": c.get("parent_content", ""),
        }
        for c in chunks
    ]
    await asyncio.to_thread(
        collection.add,
        ids=ids, embeddings=embeddings, metadatas=metadatas, documents=texts_to_embed
    )

    # 2. Store in PostgreSQL
    db_chunks = []
    for i, c in enumerate(chunks):
        chunk_obj = Chunk(
            document_id=document.id,
            content=c["text"],
            page_number=c["page_number"],
            section=c.get("section"),
            chunk_index=c["chunk_index"],
            token_count=c.get("token_count", 0),
            parent_content=c.get("parent_content"),
        )
        db.add(chunk_obj)
        db_chunks.append((chunk_obj, embeddings[i]))
        
    await db.commit() # Commit to get chunk IDs
    
    for chunk_obj, emb in db_chunks:
        emb_obj = ChunkEmbedding(
            chunk_id=chunk_obj.id,
            provider=provider_name,
            model=model_name,
            dimension=len(emb),
            embedding=emb
        )
        db.add(emb_obj)
        
    await db.commit()


async def process_document_standalone(document_id: str, ocr: bool = False, enqueue_time: float = None):
    """
    Standalone background task entry point.
    """
    if enqueue_time is None:
        enqueue_time = asyncio.get_event_loop().time()
        
    # 1. Update status to "processing" immediately so the UI shows it as processing
    from core.database import AsyncSessionLocal
    try:
        async with AsyncSessionLocal() as db:
            document = await db.get(Document, document_id)
            if document and document.status == "uploading":
                document.status = "processing"
                await db.commit()
                await broadcast_status(document, db)
    except Exception as e:
        logger.warning(f"Could not set initial processing status for document {document_id}: {e}")

    # 2. Acquire sequential ingestion semaphore and run the heavy ingestion process
    async with _ingestion_semaphore:
        async with AsyncSessionLocal() as db:
            from services.pipeline import DocumentPipeline
            pipeline = DocumentPipeline(db)
            await pipeline.process(document_id, ocr, enqueue_time)
